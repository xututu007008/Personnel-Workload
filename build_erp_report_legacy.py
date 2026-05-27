from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


BASE = Path(__file__).resolve().parent
SOURCE = BASE / "ERP使用情况-01.xlsx"
OUT_DIR = BASE / "outputs"
REPORT_STEM = "职能部门ERP系统操作分析报告"
OUT_MD = OUT_DIR / f"{REPORT_STEM}.md"
OUT_DOCX = OUT_DIR / f"{REPORT_STEM}.docx"
OUT_PDF = OUT_DIR / f"{REPORT_STEM}.pdf"
WORD_FONT = "微软雅黑"


@dataclass
class ReportData:
    source_name: str
    total_row: pd.Series
    roster_count: int
    matched_people: int
    no_ops_count: int
    source_participants: int
    ops_not_roster: list[str]
    dept_mismatch: pd.DataFrame
    dept_rows: list[list[str]]
    role_rows: list[list[str]]
    top_people_rows: list[list[str]]
    reject_top_rows: list[list[str]]
    comment_top_rows: list[list[str]]
    tier_rows: list[list[str]]
    no_ops_rows: list[list[str]]
    bullets: list[str]
    paragraphs: dict[str, str]


def clean_text(series: pd.Series) -> pd.Series:
    return series.astype(str).str.strip().replace({"nan": ""})


def fmt_int(value) -> str:
    try:
        return f"{int(round(float(value))):,}"
    except Exception:
        return ""


def pct(num, den) -> str:
    if not den:
        return "0.0%"
    return f"{num / den * 100:.1f}%"


def markdown_table(headers: list[str], rows: Iterable[Iterable[str]]) -> str:
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        escaped = [str(v).replace("\n", "<br>") for v in row]
        lines.append("| " + " | ".join(escaped) + " |")
    return "\n".join(lines)


def make_rows(df: pd.DataFrame, columns: list[str]) -> list[list[str]]:
    return [[str(row[col]) for col in columns] for _, row in df.iterrows()]


def analyze() -> ReportData:
    summary = pd.read_excel(SOURCE, sheet_name="职能部门ERP操作汇总")
    ops = pd.read_excel(SOURCE, sheet_name="职能部门人员ERP操作")
    roster = pd.read_excel(SOURCE, sheet_name="职能部门人员台账")

    summary["部门"] = clean_text(summary["部门"])
    ops["部门"] = clean_text(ops["部门"])
    ops["人员名称"] = clean_text(ops["人员名称"])
    ops["参与类型"] = clean_text(ops["参与类型"])
    roster["部门"] = clean_text(roster["部门"])
    roster["姓名"] = clean_text(roster["姓名"])
    roster["职位"] = clean_text(roster.get("职位", pd.Series([""] * len(roster))))
    roster["在职状态"] = clean_text(roster.get("在职状态", pd.Series([""] * len(roster))))

    total_row = summary[summary["部门"] == "合计"].iloc[0]
    person_ops_dept = (
        ops.groupby(["部门", "人员名称"], as_index=False)
        .agg(
            参与流程数量=("参与流程数量", "sum"),
            处理时间分钟=("处理时间（分钟）", "sum"),
            驳回次数=("驳回次数", "sum"),
            有效评论次数=("有效评论次数", "sum"),
        )
        .sort_values("参与流程数量", ascending=False)
    )
    role_pivot = (
        ops.pivot_table(
            index=["部门", "人员名称"],
            columns="参与类型",
            values="参与流程数量",
            aggfunc="sum",
            fill_value=0,
        )
        .reset_index()
    )
    for col in ["发起者", "参与审批者"]:
        if col not in role_pivot.columns:
            role_pivot[col] = 0
    person_ops_dept = person_ops_dept.merge(
        role_pivot[["部门", "人员名称", "发起者", "参与审批者"]],
        on=["部门", "人员名称"],
        how="left",
    )
    person_ops = (
        person_ops_dept.groupby("人员名称", as_index=False)
        .agg(
            ERP记录部门=("部门", lambda s: "、".join(sorted(set(s)))),
            参与流程数量=("参与流程数量", "sum"),
            处理时间分钟=("处理时间分钟", "sum"),
            驳回次数=("驳回次数", "sum"),
            有效评论次数=("有效评论次数", "sum"),
            发起者=("发起者", "sum"),
            参与审批者=("参与审批者", "sum"),
        )
        .sort_values("参与流程数量", ascending=False)
    )

    roster_unique = roster.drop_duplicates(["姓名"]).copy()
    merged = roster_unique.merge(person_ops, left_on="姓名", right_on="人员名称", how="left")
    for col in ["参与流程数量", "处理时间分钟", "驳回次数", "有效评论次数", "发起者", "参与审批者"]:
        merged[col] = merged[col].fillna(0)
    merged["是否有操作记录"] = merged["参与流程数量"].gt(0)
    merged["操作层级"] = pd.cut(
        merged["参与流程数量"],
        bins=[-1, 0, 99, 999, float("inf")],
        labels=["无记录", "低频(1-99)", "中频(100-999)", "高频(1000+)"],
    )

    ops_names = set(person_ops["人员名称"])
    roster_names = set(roster_unique["姓名"])
    ops_not_roster = sorted(ops_names - roster_names)
    dept_mismatch = merged[
        merged["是否有操作记录"]
        & merged["ERP记录部门"].notna()
        & merged.apply(lambda r: r["部门"] not in str(r["ERP记录部门"]).split("、"), axis=1)
    ][["姓名", "部门", "ERP记录部门"]].sort_values("姓名")

    dept_roster = (
        merged.groupby("部门", as_index=False)
        .agg(
            台账人数=("姓名", "count"),
            有操作人数=("是否有操作记录", "sum"),
            发起流程=("发起者", "sum"),
            审批流程=("参与审批者", "sum"),
            操作流程合计=("参与流程数量", "sum"),
            驳回次数=("驳回次数", "sum"),
            有效评论次数=("有效评论次数", "sum"),
        )
    )
    dept_roster["无操作人数"] = dept_roster["台账人数"] - dept_roster["有操作人数"]
    dept_roster["覆盖率"] = dept_roster.apply(lambda row: pct(row["有操作人数"], row["台账人数"]), axis=1)
    dept_roster = dept_roster.sort_values(["操作流程合计", "台账人数"], ascending=False)

    role_totals = ops.groupby("参与类型", as_index=False).agg(
        参与流程数量=("参与流程数量", "sum"),
        处理时间分钟=("处理时间（分钟）", "sum"),
        驳回次数=("驳回次数", "sum"),
        有效评论次数=("有效评论次数", "sum"),
    )
    total_flows = role_totals["参与流程数量"].sum()

    no_ops = merged[~merged["是否有操作记录"]].sort_values(["部门", "姓名"])
    no_ops_by_dept = (
        no_ops.groupby("部门")["姓名"]
        .apply(lambda s: "、".join(s.tolist()))
        .reset_index(name="未操作人员")
    )
    no_ops_by_dept["人数"] = no_ops_by_dept["未操作人员"].apply(lambda x: len(x.split("、")) if x else 0)
    no_ops_by_dept = no_ops_by_dept.sort_values(["人数", "部门"], ascending=[False, True])

    tier = (
        merged.groupby(["部门", "操作层级"], observed=False)["姓名"]
        .count()
        .reset_index(name="人数")
        .pivot(index="部门", columns="操作层级", values="人数")
        .fillna(0)
        .astype(int)
        .reset_index()
    )
    for col in ["无记录", "低频(1-99)", "中频(100-999)", "高频(1000+)"]:
        if col not in tier.columns:
            tier[col] = 0
    tier = tier[["部门", "无记录", "低频(1-99)", "中频(100-999)", "高频(1000+)"]].sort_values("无记录", ascending=False)

    matched_people = int(merged["是否有操作记录"].sum())
    no_ops_count = int(len(no_ops))
    source_participants = int(total_row["参与人数"])
    total_processes = total_row["发起流程总数"] + total_row["审批流程总数"]

    bullets = [
        f"流程总量为{fmt_int(total_processes)}次，其中审批流程{fmt_int(total_row['审批流程总数'])}次，占{pct(total_row['审批流程总数'], total_processes)}；发起流程{fmt_int(total_row['发起流程总数'])}次。",
        f"驳回共{fmt_int(total_row['驳回次数'])}次、有效评论{fmt_int(total_row['有效评论次数'])}次；驳回均发生在“参与审批者”角色，说明退回和意见表达主要沉淀在审批环节。",
        "资产财务部、公司领导、合同管理部、物资部、项目管理公司是流程量较高的部门，合计承担了绝大多数审批及处理压力。",
    ]
    if ops_not_roster:
        bullets.append("数据口径需关注：操作明细中有记录但人员台账未找到姓名的人员为" + "、".join(ops_not_roster) + "。")
    if not dept_mismatch.empty:
        mismatch_text = "；".join(
            f"{row['姓名']}（台账部门：{row['部门']}，ERP记录部门：{row['ERP记录部门']}）"
            for _, row in dept_mismatch.iterrows()
        )
        bullets.append(
            "另有人员姓名可匹配但部门口径不一致："
            + mismatch_text
            + "。本报告按姓名认定其“有操作记录”，部门覆盖统计按台账部门归属。"
        )

    paragraphs = {
        "overall": (
            f"近一年职能部门ERP操作呈现“覆盖面较广、审批端高度集中、评论和驳回集中于少数流程岗位”的特征。"
            f"台账共{fmt_int(len(roster_unique))}人，其中{fmt_int(matched_people)}人在台账内匹配到操作记录，"
            f"{fmt_int(no_ops_count)}人未匹配到操作记录，台账口径覆盖率为{pct(matched_people, len(roster_unique))}。"
            f"部门汇总表显示参与人数为{fmt_int(source_participants)}人，较台账姓名匹配口径多1人，差异来自操作明细中存在但台账未找到的人员。"
        ),
        "department": (
            "从覆盖率看，资产财务部、公司领导、物资部、人力资源部、技术质量部、信息化中心等部门台账成员均有操作记录；"
            "待岗人员、综合办公室、技术中心、市场开发部、项目管理公司存在较多未操作人员，应结合岗位属性判断是否属于正常无需使用ERP，或存在授权、流程分工未覆盖的情况。"
        ),
        "role": (
            "角色结构上，审批流程数量远高于发起流程数量，说明职能部门成员更多承担流程审核、把关、会签和管理反馈职能。"
            "发起者角色没有驳回记录，评论数量也明显低于审批者，符合发起端提交、审批端纠偏的业务特征。"
        ),
        "people": (
            "成员层面操作量集中度较高。袁雯雯、邢阳、李玉生、朱彬、高英昌、马言、孙艳强等人员流程量居前，"
            "反映这些岗位在合同、财务、审计、项目管理等流程链条中承担关键处理或审批节点。建议对高频人员关注授权备份、流程代理和异常积压风险，避免关键人员缺岗影响流程周转。"
        ),
        "comments": (
            "驳回次数靠前人员集中在合同管理部、资产财务部、物资部、项目管理公司和公司领导层，说明这些岗位对流程合规性、资料完整性或业务条件的校验较多。"
            "评论次数靠前人员与驳回高频人员高度重合，表明有效评论很可能承担了补充说明、修改指引和审批意见留痕作用。"
        ),
        "no_ops": (
            "无操作人员不必然代表异常，需结合岗位职责、是否待岗、是否需要ERP授权等情况综合判断。"
            "但对在职且岗位应参与流程的人员，建议由部门确认授权状态、流程节点配置及年度实际分工。"
        ),
        "source": (
            f"数据来源：{SOURCE.name}；统计口径：以“职能部门人员ERP操作”按部门、人员、参与类型汇总，"
            "并与“职能部门人员台账”按姓名匹配；姓名可匹配但部门口径不同的人员，按台账部门纳入部门覆盖统计。"
        ),
    }

    dept_rows = [
        [
            row["部门"],
            fmt_int(row["台账人数"]),
            fmt_int(row["有操作人数"]),
            fmt_int(row["无操作人数"]),
            row["覆盖率"],
            fmt_int(row["发起流程"]),
            fmt_int(row["审批流程"]),
            fmt_int(row["驳回次数"]),
            fmt_int(row["有效评论次数"]),
        ]
        for _, row in dept_roster.iterrows()
    ]
    role_rows = [
        [
            row["参与类型"],
            fmt_int(row["参与流程数量"]),
            pct(row["参与流程数量"], total_flows),
            fmt_int(row["处理时间分钟"]),
            fmt_int(row["驳回次数"]),
            fmt_int(row["有效评论次数"]),
        ]
        for _, row in role_totals.sort_values("参与流程数量", ascending=False).iterrows()
    ]
    top_people_rows = [
        [
            row["ERP记录部门"],
            row["人员名称"],
            fmt_int(row["参与流程数量"]),
            fmt_int(row["发起者"]),
            fmt_int(row["参与审批者"]),
            fmt_int(row["驳回次数"]),
            fmt_int(row["有效评论次数"]),
        ]
        for _, row in person_ops.head(15).iterrows()
    ]
    reject_top_rows = [
        [
            row["ERP记录部门"],
            row["人员名称"],
            fmt_int(row["参与流程数量"]),
            fmt_int(row["驳回次数"]),
            fmt_int(row["有效评论次数"]),
        ]
        for _, row in person_ops.sort_values("驳回次数", ascending=False).head(10).iterrows()
    ]
    comment_top_rows = [
        [
            row["ERP记录部门"],
            row["人员名称"],
            fmt_int(row["参与流程数量"]),
            fmt_int(row["驳回次数"]),
            fmt_int(row["有效评论次数"]),
        ]
        for _, row in person_ops.sort_values("有效评论次数", ascending=False).head(10).iterrows()
    ]
    tier_rows = make_rows(tier, ["部门", "无记录", "低频(1-99)", "中频(100-999)", "高频(1000+)"])
    no_ops_rows = make_rows(no_ops_by_dept, ["部门", "人数", "未操作人员"])

    return ReportData(
        source_name=SOURCE.name,
        total_row=total_row,
        roster_count=len(roster_unique),
        matched_people=matched_people,
        no_ops_count=no_ops_count,
        source_participants=source_participants,
        ops_not_roster=ops_not_roster,
        dept_mismatch=dept_mismatch,
        dept_rows=dept_rows,
        role_rows=role_rows,
        top_people_rows=top_people_rows,
        reject_top_rows=reject_top_rows,
        comment_top_rows=comment_top_rows,
        tier_rows=tier_rows,
        no_ops_rows=no_ops_rows,
        bullets=bullets,
        paragraphs=paragraphs,
    )


def write_markdown(data: ReportData) -> None:
    total_processes = data.total_row["发起流程总数"] + data.total_row["审批流程总数"]
    lines = [
        "# 职能部门成员ERP系统操作分析报告",
        "",
        f"基于《{data.source_name}》中“职能部门人员ERP操作”与“职能部门人员台账”数据。",
        "",
        "## 核心指标",
        "",
        markdown_table(
            ["台账人员", "有操作记录", "无操作记录", "操作覆盖率", "流程总量"],
            [[
                fmt_int(data.roster_count),
                fmt_int(data.matched_people),
                fmt_int(data.no_ops_count),
                pct(data.matched_people, data.roster_count),
                fmt_int(total_processes),
            ]],
        ),
        "",
        "## 一、总体结论",
        "",
        data.paragraphs["overall"],
        "",
    ]
    lines.extend(f"- {item}" for item in data.bullets)
    lines.extend([
        "",
        "## 二、部门覆盖与操作量",
        "",
        markdown_table(["部门", "台账人数", "有操作", "无操作", "覆盖率", "发起", "审批", "驳回", "评论"], data.dept_rows),
        "",
        data.paragraphs["department"],
        "",
        "## 三、角色结构分析",
        "",
        markdown_table(["角色", "流程数量", "占比", "处理时间(分钟)", "驳回", "有效评论"], data.role_rows),
        "",
        data.paragraphs["role"],
        "",
        "## 四、成员操作强度",
        "",
        markdown_table(["部门", "姓名", "流程合计", "发起", "审批", "驳回", "有效评论"], data.top_people_rows),
        "",
        data.paragraphs["people"],
        "",
        "## 五、评论与驳回情况",
        "",
        "### 驳回次数前十",
        "",
        markdown_table(["部门", "姓名", "流程合计", "驳回次数", "有效评论"], data.reject_top_rows),
        "",
        data.paragraphs["comments"],
        "",
        "### 有效评论前十",
        "",
        markdown_table(["部门", "姓名", "流程合计", "驳回次数", "有效评论"], data.comment_top_rows),
        "",
        "## 六、未操作人员与分层",
        "",
        markdown_table(["部门", "无记录", "低频(1-99)", "中频(100-999)", "高频(1000+)"], data.tier_rows),
        "",
        data.paragraphs["no_ops"],
        "",
        markdown_table(["部门", "未操作人数", "未操作人员"], data.no_ops_rows),
        "",
        "## 七、管理建议",
        "",
        "- 按部门复核无操作人员名单，区分“无需使用ERP”“授权未开通”“流程节点未配置”“岗位变动/离职待更新”等原因。",
        "- 对高频审批人员建立代理和备岗机制，重点关注资产财务部、合同管理部、物资部、项目管理公司等高负荷部门。",
        "- 针对驳回和评论高发环节，梳理常见退回原因，形成发起端资料清单或审批前校验规则，减少重复修改。",
        "- 统一操作明细与人员台账口径，补齐操作明细中存在但台账未找到的人员信息，并复核姓名匹配但部门口径不一致的人员，避免后续考核或授权分析出现偏差。",
        "",
        data.paragraphs["source"],
        "",
    ])
    OUT_MD.write_text("\n".join(lines), encoding="utf-8")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = WORD_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT)
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = WORD_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Title", 22, "0B2545"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = WORD_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)


def add_doc_table(doc: Document, headers: list[str], rows: Iterable[Iterable], widths_cm: list[float] | None = None) -> None:
    rows = list(rows)
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.cell(0, i), header, bold=True, color="0B2545")
        set_cell_shading(table.cell(0, i), "F2F4F7")
    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row):
            set_cell_text(table.cell(row_index, col_index), value)
    if widths_cm:
        for row in table.rows:
            for i, width in enumerate(widths_cm):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()


def add_kv_table(doc: Document, items: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for i, (label, value) in enumerate(items):
        cell = table.cell(0, i)
        set_cell_shading(cell, "E8EEF5")
        cell.text = ""
        p1 = cell.paragraphs[0]
        r1 = p1.add_run(label)
        r1.font.name = WORD_FONT
        r1._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT)
        r1.font.size = Pt(8.5)
        r1.font.color.rgb = RGBColor.from_string("44546A")
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = cell.add_paragraph()
        r2 = p2.add_run(value)
        r2.bold = True
        r2.font.name = WORD_FONT
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT)
        r2.font.size = Pt(15)
        r2.font.color.rgb = RGBColor.from_string("0B2545")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = WORD_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT)
    run.font.size = Pt(10.5)


def write_docx(data: ReportData) -> None:
    total_processes = data.total_row["发起流程总数"] + data.total_row["审批流程总数"]
    doc = Document()
    set_doc_defaults(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("职能部门成员ERP系统操作分析报告")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"基于《{data.source_name}》中“职能部门人员ERP操作”与“职能部门人员台账”数据")
    run.font.name = WORD_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("44546A")

    add_kv_table(
        doc,
        [
            ("台账人员", fmt_int(data.roster_count)),
            ("有操作记录", fmt_int(data.matched_people)),
            ("无操作记录", fmt_int(data.no_ops_count)),
            ("操作覆盖率", pct(data.matched_people, data.roster_count)),
            ("流程总量", fmt_int(total_processes)),
        ],
    )
    doc.add_heading("一、总体结论", level=1)
    doc.add_paragraph(data.paragraphs["overall"])
    for item in data.bullets:
        add_bullet(doc, item)

    doc.add_heading("二、部门覆盖与操作量", level=1)
    add_doc_table(doc, ["部门", "台账人数", "有操作", "无操作", "覆盖率", "发起", "审批", "驳回", "评论"], data.dept_rows, [2.4, 1.3, 1.3, 1.3, 1.4, 1.3, 1.5, 1.3, 1.3])
    doc.add_paragraph(data.paragraphs["department"])

    doc.add_heading("三、角色结构分析", level=1)
    add_doc_table(doc, ["角色", "流程数量", "占比", "处理时间(分钟)", "驳回", "有效评论"], data.role_rows, [2.4, 2, 1.5, 2.4, 1.5, 1.8])
    doc.add_paragraph(data.paragraphs["role"])

    doc.add_heading("四、成员操作强度", level=1)
    add_doc_table(doc, ["部门", "姓名", "流程合计", "发起", "审批", "驳回", "有效评论"], data.top_people_rows, [2.4, 1.7, 1.7, 1.4, 1.7, 1.4, 1.7])
    doc.add_paragraph(data.paragraphs["people"])

    doc.add_heading("五、评论与驳回情况", level=1)
    doc.add_heading("驳回次数前十", level=2)
    add_doc_table(doc, ["部门", "姓名", "流程合计", "驳回次数", "有效评论"], data.reject_top_rows, [2.5, 1.8, 2, 1.8, 1.8])
    doc.add_paragraph(data.paragraphs["comments"])
    doc.add_heading("有效评论前十", level=2)
    add_doc_table(doc, ["部门", "姓名", "流程合计", "驳回次数", "有效评论"], data.comment_top_rows, [2.5, 1.8, 2, 1.8, 1.8])

    doc.add_heading("六、未操作人员与分层", level=1)
    add_doc_table(doc, ["部门", "无记录", "低频(1-99)", "中频(100-999)", "高频(1000+)"], data.tier_rows, [3, 1.6, 2, 2, 2])
    doc.add_paragraph(data.paragraphs["no_ops"])
    add_doc_table(doc, ["部门", "未操作人数", "未操作人员"], data.no_ops_rows, [2.4, 1.8, 10.2])

    doc.add_heading("七、管理建议", level=1)
    for item in [
        "按部门复核无操作人员名单，区分“无需使用ERP”“授权未开通”“流程节点未配置”“岗位变动/离职待更新”等原因。",
        "对高频审批人员建立代理和备岗机制，重点关注资产财务部、合同管理部、物资部、项目管理公司等高负荷部门。",
        "针对驳回和评论高发环节，梳理常见退回原因，形成发起端资料清单或审批前校验规则，减少重复修改。",
        "统一操作明细与人员台账口径，补齐操作明细中存在但台账未找到的人员信息，并复核姓名匹配但部门口径不一致的人员，避免后续考核或授权分析出现偏差。",
    ]:
        add_bullet(doc, item)
    doc.add_paragraph(data.paragraphs["source"])
    doc.save(OUT_DOCX)


def resolve_pdf_font() -> str:
    candidates = [
        (WORD_FONT, Path("/Library/Fonts/Microsoft YaHei.ttf")),
        (WORD_FONT, Path("/System/Library/Fonts/Supplemental/Microsoft YaHei.ttf")),
        ("ArialUnicode", Path("/Library/Fonts/Arial Unicode.ttf")),
        ("ArialUnicode", Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf")),
        ("STHeiti", Path("/System/Library/Fonts/STHeiti Medium.ttc")),
        ("Songti", Path("/System/Library/Fonts/Supplemental/Songti.ttc")),
    ]
    for name, path in candidates:
        if path.exists():
            try:
                pdfmetrics.registerFont(TTFont(name, str(path)))
                return name
            except Exception:
                continue
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    return "STSong-Light"


def pdf_paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(str(text).replace("\n", "<br/>"), style)


def pdf_table(headers: list[str], rows: list[list[str]], style: ParagraphStyle, col_widths: list[float] | None = None) -> Table:
    table_data = [[pdf_paragraph(h, style) for h in headers]]
    table_data.extend([[pdf_paragraph(value, style) for value in row] for row in rows])
    table = Table(table_data, colWidths=col_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
                ("FONTNAME", (0, 0), (-1, -1), style.fontName),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D0D5DD")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def write_pdf(data: ReportData) -> None:
    font = resolve_pdf_font()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ChineseTitle", parent=styles["Title"], fontName=font, fontSize=20, leading=26, alignment=TA_CENTER, textColor=colors.HexColor("#0B2545"), spaceAfter=8)
    subtitle_style = ParagraphStyle("ChineseSubtitle", parent=styles["Normal"], fontName=font, fontSize=9, leading=13, alignment=TA_CENTER, textColor=colors.HexColor("#44546A"), spaceAfter=12)
    h1 = ParagraphStyle("ChineseH1", parent=styles["Heading1"], fontName=font, fontSize=14, leading=18, textColor=colors.HexColor("#2E74B5"), spaceBefore=10, spaceAfter=6)
    h2 = ParagraphStyle("ChineseH2", parent=styles["Heading2"], fontName=font, fontSize=11, leading=15, textColor=colors.HexColor("#1F4D78"), spaceBefore=8, spaceAfter=4)
    body = ParagraphStyle("ChineseBody", parent=styles["BodyText"], fontName=font, fontSize=9, leading=14, alignment=TA_LEFT, spaceAfter=6)
    small = ParagraphStyle("ChineseSmall", parent=body, fontSize=7, leading=10)

    doc = SimpleDocTemplate(
        str(OUT_PDF),
        pagesize=landscape(A4),
        leftMargin=1.2 * cm,
        rightMargin=1.2 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
    )
    total_processes = data.total_row["发起流程总数"] + data.total_row["审批流程总数"]
    story = [
        pdf_paragraph("职能部门成员ERP系统操作分析报告", title_style),
        pdf_paragraph(f"基于《{data.source_name}》中“职能部门人员ERP操作”与“职能部门人员台账”数据", subtitle_style),
        pdf_table(
            ["台账人员", "有操作记录", "无操作记录", "操作覆盖率", "流程总量"],
            [[fmt_int(data.roster_count), fmt_int(data.matched_people), fmt_int(data.no_ops_count), pct(data.matched_people, data.roster_count), fmt_int(total_processes)]],
            body,
            [3.2 * cm, 3.2 * cm, 3.2 * cm, 3.2 * cm, 3.2 * cm],
        ),
        Spacer(1, 0.2 * cm),
        pdf_paragraph("一、总体结论", h1),
        pdf_paragraph(data.paragraphs["overall"], body),
    ]
    story.extend(pdf_paragraph(f"• {item}", body) for item in data.bullets)
    sections = [
        ("二、部门覆盖与操作量", data.paragraphs["department"], ["部门", "台账人数", "有操作", "无操作", "覆盖率", "发起", "审批", "驳回", "评论"], data.dept_rows),
        ("三、角色结构分析", data.paragraphs["role"], ["角色", "流程数量", "占比", "处理时间(分钟)", "驳回", "有效评论"], data.role_rows),
        ("四、成员操作强度", data.paragraphs["people"], ["部门", "姓名", "流程合计", "发起", "审批", "驳回", "有效评论"], data.top_people_rows),
    ]
    for title, paragraph, headers, rows in sections:
        story.extend([pdf_paragraph(title, h1), pdf_table(headers, rows, small), pdf_paragraph(paragraph, body)])

    story.extend([
        pdf_paragraph("五、评论与驳回情况", h1),
        pdf_paragraph("驳回次数前十", h2),
        pdf_table(["部门", "姓名", "流程合计", "驳回次数", "有效评论"], data.reject_top_rows, small),
        pdf_paragraph(data.paragraphs["comments"], body),
        pdf_paragraph("有效评论前十", h2),
        pdf_table(["部门", "姓名", "流程合计", "驳回次数", "有效评论"], data.comment_top_rows, small),
        pdf_paragraph("六、未操作人员与分层", h1),
        pdf_table(["部门", "无记录", "低频(1-99)", "中频(100-999)", "高频(1000+)"], data.tier_rows, small),
        pdf_paragraph(data.paragraphs["no_ops"], body),
        pdf_table(["部门", "未操作人数", "未操作人员"], data.no_ops_rows, small),
        pdf_paragraph("七、管理建议", h1),
    ])
    for item in [
        "按部门复核无操作人员名单，区分“无需使用ERP”“授权未开通”“流程节点未配置”“岗位变动/离职待更新”等原因。",
        "对高频审批人员建立代理和备岗机制，重点关注资产财务部、合同管理部、物资部、项目管理公司等高负荷部门。",
        "针对驳回和评论高发环节，梳理常见退回原因，形成发起端资料清单或审批前校验规则，减少重复修改。",
        "统一操作明细与人员台账口径，补齐操作明细中存在但台账未找到的人员信息，并复核姓名匹配但部门口径不一致的人员，避免后续考核或授权分析出现偏差。",
    ]:
        story.append(pdf_paragraph(f"• {item}", body))
    story.append(pdf_paragraph(data.paragraphs["source"], body))
    doc.build(story)


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    data = analyze()
    write_markdown(data)
    write_docx(data)
    write_pdf(data)
    print(OUT_MD)
    print(OUT_DOCX)
    print(OUT_PDF)


if __name__ == "__main__":
    main()
